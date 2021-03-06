import os
import re
import click

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches
from docx.shared import Pt

from utils.utils_print import warningPrint


def make_document(vtt_file_path, title, sentence_count):
    document = Document()

    style_document_default_vtt(document, vtt_file_path, title)

    table = make_formatted_table(document)
    vtt_lines = parse_vtt_file(vtt_file_path, sentence_count)

    fill_table(table, vtt_lines)
    set_col_widths(table)

    # This is super jank, I'm aware. Plan to make more robust. (Currently chops
    # off ./input at the beginning and .vtt at the end)
    output_name = vtt_file_path[7:-4]
    document.save(f'./output/{output_name}.docx')

def style_document_default_vtt(document, vtt_file_path, title):
    style = document.styles['Normal']
    font = style.font
    font.name = 'Lora'
    font.size = Pt(12)
    heading = document.add_paragraph()
    run = heading.add_run(f'Transcript - "{title}" video')
    run.bold = True
    font = run.font
    font.name = 'Montserrat'
    font.size = Pt(11)

def fill_table(table, vtt_lines):
    timeReg = re.compile(r'\d{2}:\d{2}:\d{2}')
    count = 0
    new_row = table.add_row()
    for line in vtt_lines:
        if timeReg.search(line):
            timestamps = line.split(' --> ')
            new_row.cells[0].text = timestamps[0][:-4] 
            count = count + 1

        if not timeReg.search(line):
            new_row.cells[2].text = line+'\n'
            count = count + 1

        if count >= 3:
            new_row = table.add_row()
            count = 0


def make_formatted_table(document):
    # The repetition in here is necessary because ms word demands it
    table = document.add_table(rows=1, cols=3)
    top_row = table.rows[0]
    heading_cells = top_row.cells
    run = heading_cells[0].paragraphs[0].add_run('Timecode\n')
    run.bold = True
    run = heading_cells[1].paragraphs[0].add_run('Speaker\n')
    run.bold = True
    run = heading_cells[2].paragraphs[0].add_run('Dialogue\n')
    run.bold = True
    table.style = 'Table Grid'
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="d9d9d9"/>'.format(nsdecls('w')))
    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="d9d9d9"/>'.format(nsdecls('w')))
    shading_elm_3 = parse_xml(r'<w:shd {} w:fill="d9d9d9"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)
    table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_3)

    return table


def get_vtt_files(directory_name):
    list_of_file = os.listdir(directory_name)
    all_files = []
    vtt_reg = re.compile(r'.vtt$')

    for entry in list_of_file:
        full_path = os.path.join(directory_name, entry)

        if os.path.isdir(full_path):
            all_files = all_files + get_vtt_files(full_path)
            continue

        if vtt_reg.search(full_path):
            all_files.append(full_path)

    return all_files


# Removes unnecessary vtt stuff that shouldn't be in the output document
# includes: WEBVTT tag, line numbers, any undesired_lines passed in
def remove_unwanted_content(file_lines, undesired_lines):
    undesired_lines = undesired_lines + ['WEBVTT\n'] + ['\n']
    desired_lines = []
    num_reg = re.compile(r'^\d{1,4}\n')
    for line in file_lines:
        if line in undesired_lines:
            continue
        if num_reg.search(line):
            continue
        desired_lines.append(line)

    return desired_lines


def read_file_lines(vtt_file):
    with open(vtt_file, 'r') as f:
        lines = f.readlines()

    return lines


def parse_vtt_file(vtt_file, sentence_count):
    count = 0
    end_reg = re.compile(r'[a-z](\.|!|\?)')
    full_line = ''
    output_lines = []
    time_recording = True
    time_reg = re.compile(r'\d{1,2}:\d{1,2}:\d{1,2}')
    lines = read_file_lines(vtt_file)
    desired_lines = remove_unwanted_content(lines, [])

    # combine sentences
    for line in desired_lines:
        if end_reg.search(line):
            count = count + 1
            if count >= sentence_count:
                full_line += line
                full_line = full_line.replace('\n', ' ')
                output_lines.append(full_line)
                full_line = ''
                time_recording = True
                count = 0
                continue

        if time_recording and time_reg.search(line):
            time_recording = False 
            output_lines.append(line)
            continue

        if not time_reg.search(line):
            full_line += line

    full_line = full_line.replace('\n', ' ')
    output_lines.append(full_line)
    print(output_lines)

    return output_lines


def set_col_widths(table):
    widths = (Inches(1), Inches(1), Inches(4))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

